import os
import sys
import docx

def convert_to_writable(kanji_dict):
    """
    Returns:
        {
            '亜/亞': {
                'ア': '亜流 亜麻 亜熱帯'
            },
            '哀': {
                'アイ': '哀愁 哀願 悲哀',
                'あわれ': '哀れ 哀れな話 哀れがる',
                'あわれむ': '哀れむ 哀れみ'
            },
        ...
        }
    """
    # Initialize an empty dictionary to store the converted kanji data.
    writable = {}

    # Define a mapping of numerical keys to specific symbols.
    d = {'1': '', '2': '+', '3': '*', '4': '^'}
    
    # Iterate over each item in the input dictionary (kanji_dict).
    for a, b in kanji_dict.items():
        # For each 'kanji' item, transform the tuples in the list b['kanji']
        # using the mapping defined in 'd'. Each tuple consists of a number
        # (indicating the type) and a string (the actual character or phonetic).
        # The transformation appends the mapped symbol (from dictionary 'd')
        # to the phonetic/stroke count part of the tuple.
        # The transformed items are joined into a single string with slashes '/' as separators.
        key = '/'.join(list(map(lambda x: x[1] + d[x[0]], b['kanji'])))

        # Map the transformed string (key) to the corresponding 'kana' value from
        # the input dictionary, adding it to the 'writable' dictionary.
        writable[key] = b['kana']

    # Return the transformed/writable dictionary.
    return writable


def prepare_directory(filepath):
    directory = os.path.dirname(filepath)

    # Check if directory exists, create if it doesn't
    if not os.path.exists(directory):
        os.makedirs(directory)
        print(f"Created directory: {directory}")
 

def prepare_output_path(filepath):
    prepare_directory(filepath)

    filename = os.path.basename(filepath)
    # Check if file exists
    if os.path.exists(filepath):
        while True:
            user_input = input(f"File {filename} already exists. Delete and continue? (y/n): ").lower()
            if user_input == 'y':
                os.remove(filepath)
                print(f"Deleted existing file: {filepath}")
                break
            elif user_input == 'n':
                print("Exiting program.")
                sys.exit(0)
            else:
                print("Invalid input. Please enter 'y' or 'n'.")
    
    return filepath


def readfile(dirname):
    data = []
    for f in os.listdir(dirname):
        fpath = os.path.join(dirname, f)
        if not os.path.isfile(fpath):
            continue
        h = open(fpath, 'r')
        for i in h.readlines():
            if i.strip() == '':
                continue
            data.append(i.strip())
        h.close()
    return data


def save_to_docx(sets, filename, header, filepath):
    # Import the Document class only once at the function start
    doc = docx.Document()
    doc.add_heading(header, level=1)
    
    # Declaring constants for font sizes to avoid repetitive calls and easy updates
    yomi_font_size = docx.shared.Pt(16)
    kanji_font_size = docx.shared.Pt(10)

    # Function to set paragraph format. Encapsulates the repeated style assignment.
    def set_paragraph_format(paragraph):
        pf = paragraph.paragraph_format
        pf.left_indent = docx.shared.Inches(-1)
        pf.right_indent = docx.shared.Inches(-1)
        pf.line_spacing = docx.shared.Pt(0)
        pf.space_before = docx.shared.Pt(0)
        pf.space_after = docx.shared.Pt(0)

    # Iterate over sorted keys of the sets dictionary
    for a in sorted(sets.keys()):
        p = doc.add_paragraph()
        set_paragraph_format(p)

        # Adding yomi with bold and specific font size
        yomi = p.add_run(a + '\t')
        yomi.bold = True
        yomi.font.size = yomi_font_size

        # Adding kanji with non-bold and specific font size
        kanji = p.add_run('、'.join(sets[a].keys()))
        kanji.bold = False
        kanji.font.size = kanji_font_size

    # Construct a full file path with filename extension explicitly as ".docx"
    full_path = os.path.join(filepath, f"{filename}")
    doc.save(full_path)


def save_to_docx_old(sets, filename, filepath):
    doc = docx.Document()
    doc.add_heading(filename)
    
    p = doc.add_table(rows=1, cols=2)
    p.style.paragraph_format.left_indent = docx.shared.Inches(-1)
    p.rows[0].cells[0].text = name
    p.rows[0].cells[1].text = '漢字'
    for a in sorted(sets.keys()):
        r = p.add_row().cells
        r[0].text = a
        r[1].text = '、'.join(sets[a].keys())

    doc.save('%s/%s.docx' %(filepath, filename))


def save_to_csv(sets, filepath):
    filepath = prepare_output_path(filepath)

    with open(filepath, 'w') as h:
        h.write('漢字, 読み\n')
        for a in sorted(sets.keys()):
            yomi = '、'.join(sets[a].keys())
            h.write(f'{a}, {yomi}\n')

